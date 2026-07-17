import os
import time
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel,Field
load_dotenv()
my_api_key=os.getenv("GROQ_API_KEY")
if not my_api_key:
    raise ValueError("There is no Api key")
client=Groq(api_key=my_api_key)
model="openai/gpt-oss-120b"
job_description="""
Description
"This role is intended for 2026 and 2025 graduates only."

At Amazon, we hire the best minds in technology to innovate and build on behalf of our customers. The focus we have on our customers is why we are one of the world’s most beloved brands – customer obsession is part of our company DNA. Our Software Development Engineers (SDEs) use cutting-edge technology to solve complex problems and get to see the impact of their work first-hand. The challenges SDEs solve for at Amazon are big and influence millions of customers, sellers, and products around the world. We are looking for individuals who are passionate about creating new products, features, and services from scratch while managing ambiguity and the pace of a company where development cycles are measured in weeks, not years. If this sounds interesting to you, apply and come chart your own path at Amazon.
Applications are reviewed on a rolling basis. For an update on your status, or to confirm your application was submitted successfully, please login to your candidate portal. NOTE: Amazon works with a high volume of applicants, so we appreciate your patience as we review applications

Key job responsibilities
• Collaborate with experienced cross-disciplinary Amazonians to conceive, design, and bring innovative products and services to market.
• Design and build innovative technologies in a large distributed computing environment and help lead fundamental changes in the industry.
• Create solutions to run predictions on distributed systems with exposure to innovative technologies at incredible scale and speed.
• Build distributed storage, index, and query systems that are scalable, fault-tolerant, low cost, and easy to manage/use.
• Design and code the right solutions starting with broadly defined problems.
• Work in an agile environment to deliver high-quality software.

Basic Qualifications
- Bachelor's degree or above in computer science, computer engineering, or related field
- Knowledge of Computer Science fundamentals such as object-oriented design, algorithm design, data structures, problem solving, and complexity analysis.
- Knowledge of programming languages such as C/C++, Python, Java or Perl

Preferred Qualifications
- Previous technical internship(s).
- Experience with distributed, multi-tiered systems, algorithms, and relational databases.
- Experience in optimization mathematics such as linear programming and nonlinear optimization.
- Effectively articulate technical challenges and solutions.
- Adept at handling ambiguous or undefined problems as well as ability to think abstractly.
"""
class JobD(BaseModel):
    role:str
    required_skills:list[str]
    prefered_skills:list[str]
    minimum_experiences:float | None
    education_requirments:list[str]
    responsibilities:list[str]

jobd_schema=JobD.model_json_schema()

system_prompt=f"""
You are an expert HR assistant.

Your job is to analyze job descriptions and extract
structured information from them.

Return ONLY valid JSON matching this schema:

{jobd_schema}
IMPORTANT:
Do NOT return the schema itself.
Do NOT return fields like "properties", "title" or "type".
Fill the schema with actual information extracted from the job description.

If minimum experience is not mentioned, return null.
If information for a list is missing, return an empty list.
Do not invent information.
"""
user_prompt=f"""
Analyze the following Job Description
{job_description}
"""
system_message={
    "role":"system",
    "content":system_prompt
}
message_user={
    "role":"user",
    "content":user_prompt
}
response_format={
    "type":"json_object"
}
messages=[system_message,message_user]
response=client.chat.completions.create(model=model,messages=messages,response_format=response_format)

answer=response.choices[0].message.content
raw_json=answer

import json
job_data=json.loads(raw_json)

job=JobD(**job_data)
print(job.minimum_experiences)
print(job.education_requirments)

#part 2
class MatchResult(BaseModel):
    score:float
    details:dict
class Experience(BaseModel):
    company: str | None = None
    role:str | None = None
    duration:str | None = None
    description:str | None = None
    skills_used:list[str]=[]
class Resume(BaseModel):
    name:str | None = None
    email:str | None = None
    phone:str | None = None
    total_experience_years: float | None = None
    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []

resume_schema=Resume.model_json_schema()
# part 3 match result and parse resume function:
def final_score(job,resume):
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    You are an HR recruiter.

    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    Return JSON matching this schema:

    {match_schema}

    Give me:

    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.
    """
    message={
        "role": "user",
        "content" : prompt
    }
    messages=[message]
    response_format={
        "type": "json_object"
    }
    response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)
def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt = f"""
    Parse the following resume:

    {resume_text}
    """
    message_system={
        "role" : "system",
        "content" : system_prompt
    }
    message_user={
        "role" : "user",
        "content" : user_prompt
    }
    messages=[message_system, message_user]
    response_format={
        "type": "json_object"
    }
    response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume

from pypdf import PdfReader
from docx import Document
def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        return None

#Part 4
resume_folder = Path("resumes")
all_results=[]
for file_path in resume_folder.iterdir():
    #C:\Users\Pratyush\padho_with_pratyush\week1\day5\resumes\abhay resume new - Abhay Singh.pdf
    if file_path.suffix.lower() not in [".pdf", ".docx"]:
        continue
    print("\nProcessing:", file_path.name)
    resume_text = read_resume(file_path)
    parsed_resume=parse_resume(resume_text) # llm call1
    time.sleep(5)
    result = final_score(job, parsed_resume) #llm caLL2
    #score and details
    #Taking sleep so that it doesnt hang the server
    time.sleep(5)
    print("Score:", result.score)
    all_results.append({
        "name": parsed_resume.name,
        "score": result.score,
        "details": result.details
    })
all_results.sort(
    key=lambda candidate: candidate["score"],
    reverse=True
)
top_2 = all_results[:2]
worst_2 = all_results[-2:]


print("TOP 2 CANDIDATES")
for candidate in top_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )

    print(candidate["details"])

print("LOWEST 2 CANDIDATES")
for candidate in worst_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )
    print(candidate["details"])